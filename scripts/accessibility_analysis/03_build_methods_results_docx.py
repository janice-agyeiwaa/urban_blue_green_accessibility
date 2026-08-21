"""Build the private Paper 3 Methods and Results 2.0 Word deliverable.

The document is intentionally written to ``deliverables/``, which is ignored by
Git. Public evidence remains in ``outputs/accessibility_analysis``. The script reads only the
canonical Paper 3 data and generated tables/figures; it does not rerun models.
"""

from __future__ import annotations

import csv
import math
import sys
from datetime import date
from pathlib import Path


ROOT = Path(__file__).resolve().parents[2]
LOCAL_PACKAGES = ROOT / "artifacts" / "python_packages"
if LOCAL_PACKAGES.exists():
    sys.path.insert(0, str(LOCAL_PACKAGES))

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DATA_FILE = ROOT / "data" / "processed" / "accessibility_analysis" / "accessibility_analysis_by_catchment.csv"
TABLE_DIR = ROOT / "outputs" / "accessibility_analysis" / "tables"
FIGURE_DIR = ROOT / "outputs" / "accessibility_analysis" / "figures"
SOURCE_MEDIA = ROOT / "artifacts" / "docx_source_media" / "current"
DELIVERABLE_DIR = ROOT / "deliverables"
OUTPUT_FILE = DELIVERABLE_DIR / "Paper_3_Methods_and_Results_2.0.docx"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EDF4FA"
MID_GREY = "666666"
LIGHT_GREY = "F2F2F2"
WHITE = "FFFFFF"
TERM_ORDER = [
    "z_shelter",
    "z_low_income",
    "z_bachelors",
    "z_visible_minority",
    "z_indigenous",
]
OUTCOME_ORDER = ["pedestrian", "transit", "physical", "visual", "haptic", "multidimensional"]
OUTCOME_LABEL = {
    "pedestrian": "Pedestrian",
    "transit": "Transit",
    "physical": "Combined physical",
    "visual": "Visual",
    "haptic": "Haptic",
    "multidimensional": "Multidimensional",
}


def requested_output_file() -> Path:
    """Return an optional --output path, resolved from the project root."""
    prefix = "--output="
    value = next((arg[len(prefix):] for arg in sys.argv[1:] if arg.startswith(prefix)), None)
    if not value:
        return OUTPUT_FILE
    path = Path(value)
    return path if path.is_absolute() else ROOT / path


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def as_float(value: str | None) -> float | None:
    try:
        result = float(value)  # type: ignore[arg-type]
    except (TypeError, ValueError):
        return None
    return None if math.isnan(result) else result


def fmt_num(value: str | float | None, digits: int = 3) -> str:
    number = as_float(str(value)) if value is not None else None
    return "—" if number is None else f"{number:.{digits}f}"


def fmt_p(value: str | float | None) -> str:
    number = as_float(str(value)) if value is not None else None
    if number is None:
        return "—"
    if number < 0.001:
        return "< .001"
    return f"{number:.3f}".lstrip("0")


def fmt_p_statement(value: str | float | None) -> str:
    rendered = fmt_p(value)
    return f"p {rendered}" if rendered.startswith("<") else f"p = {rendered}"


def beta_from_cell(value: str) -> str:
    """Extract the coefficient from a preformatted ``beta (SE)`` table cell."""
    return value.split(" ", maxsplit=1)[0].replace("*", "")


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 60, start: int = 70, bottom: int = 60, end: int = 70) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def set_repeat_heading(paragraph, keep_with_next: bool = True) -> None:
    properties = paragraph._p.get_or_add_pPr()
    if keep_with_next:
        properties.append(OxmlElement("w:keepNext"))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor(40, 40, 40)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, colour in (
        ("Title", 24, BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, "2F5597"),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        add_page_number(footer)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    set_repeat_heading(paragraph)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MID_GREY)


def add_note(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, 90, 110, 90, 110)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_picture(document: Document, path: Path, width: float, caption: str) -> None:
    if not path.exists():
        paragraph = document.add_paragraph(f"[Figure unavailable: {path.name}]")
        paragraph.runs[0].italic = True
        add_caption(document, caption)
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(document, caption)


def add_figure_placeholder(document: Document, heading: str, instructions: str, caption: str) -> None:
    """Add a visible, editable placeholder for a manually prepared figure."""
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    row = table.rows[0]
    row.height = Inches(3.5)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, 180, 220, 180, 220)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = paragraph.add_run(heading + "\n")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor.from_string(BLUE)
    detail_run = paragraph.add_run(instructions)
    detail_run.font.size = Pt(10)
    detail_run.font.color.rgb = RGBColor.from_string(MID_GREY)
    add_caption(document, caption)


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    caption: str,
    font_size: float = 8.0,
    first_column_shading: bool = False,
) -> None:
    caption_paragraph = document.add_paragraph()
    caption_paragraph.paragraph_format.keep_with_next = True
    caption_paragraph.paragraph_format.space_after = Pt(3)
    run = caption_paragraph.add_run(caption)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for cell_run in paragraph.runs:
                cell_run.bold = True
                cell_run.font.color.rgb = RGBColor.from_string(WHITE)
                cell_run.font.size = Pt(font_size)

    for row_number, values in enumerate(rows):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if row_number % 2 == 1:
                set_cell_shading(cell, LIGHT_GREY)
            elif first_column_shading and index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for cell_run in paragraph.runs:
                    cell_run.font.size = Pt(font_size)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)


def result_lookup(rows: list[dict[str, str]], outcome: str, term: str) -> dict[str, str]:
    return next(row for row in rows if row["outcome"] == outcome and row["term"] == term)


def site_control_lookup(rows: list[dict[str, str]], outcome: str, term: str) -> dict[str, str]:
    return next(row for row in rows if row["outcome"] == outcome and row["term"] == term)


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(85)
    run = paragraph.add_run("PAPER 3")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string("5B9BD5")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Methods and Results 2.0")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(27)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Multidimensional accessibility to blue–green waterfront sites\n").bold = True
    subtitle.add_run("Metro Vancouver | 10-minute main specification")

    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(30)
    run = line.add_run("WORKING MANUSCRIPT DRAFT")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(WHITE)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), BLUE)
    line._p.get_or_add_pPr().append(shading)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(25)
    meta.add_run(f"Prepared {date.today():%B %d, %Y}\n")
    meta.add_run("Canonical Paper 3 analysis; private deliverable, not for GitHub").italic = True

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_executive_summary(document: Document, sample_rows: list[dict[str, str]]) -> None:
    add_heading(document, "What changed in this version", 1)
    document.add_paragraph(
        "This version replaces the earlier 20-minute presentation with a conceptually selected "
        "10-minute main specification and applies one reproducible quality-control and spatial-model "
        "procedure to every outcome and catchment. It retains the multidimensional outcome required "
        "by the approved proposal because that score is the accessibility input for Paper 4."
    )
    add_bullets(
        document,
        [
            "The valid 10-minute regression sample is 105 sites; three records with critical Census-quality flags are excluded consistently.",
            "All models use row-standardized seven-nearest-neighbour spatial weights, the smallest common k that connects every catchment sample.",
            "Baseline and controlled results are reported for pedestrian, transit, combined physical, visual, haptic, and multidimensional accessibility.",
            "Controlled model families are rechecked; the final 10-minute multidimensional model is SAR error because controlled OLS retained residual spatial autocorrelation.",
            "Asim’s requested pairwise correlations, median reference lines, quadrant classifications, divergent cases, spider profiles, and a ternary composition plot are regenerated from the current 114-site data.",
            "Five-, 20-, and 30-minute models are retained only as catchment sensitivity analyses.",
        ],
    )
    add_note(
        document,
        "Interpretation boundary: the results are cross-sectional associations. Positive low-income "
        "coefficients for physical access primarily reflect nearby transit-stop supply; they do not "
        "show that lower-income neighbourhoods have better park quality, amenities, safety, visual "
        "access, or haptic access.",
    )

    rows = [
        [row["walktime_min"], row["total_sites"], row["HasData_equals_1"], row["critical_flags"], row["final_regression_N"]]
        for row in sample_rows
    ]
    add_table(
        document,
        ["Walk time (min)", "Sites", "With Census data", "Critical flags", "Final N"],
        rows,
        "Summary table. Analytical sample by walking catchment",
        font_size=8.5,
    )
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_methods(
    document: Document,
    analysis_rows: list[dict[str, str]],
) -> None:
    ten_min_all = [row for row in analysis_rows if row["walktime_min"] == "10"]
    municipalities = sorted({row["municipality"] for row in ten_min_all if row["municipality"]})
    add_heading(document, "4.2 Methods", 1)
    add_heading(document, "4.2.1 Study area and analytical sites", 2)
    document.add_paragraph(
        f"The cross-sectional ecological study includes 114 publicly accessible blue–green waterfront "
        f"sites distributed across {len(municipalities)} Metro Vancouver municipalities or local "
        "jurisdictions. The site is the unit used to construct accessibility outcomes; neighbourhood "
        "socioeconomic characteristics are summarized within walking catchments surrounding each site. "
        "The current sample supersedes the 95/96-site count in the approved proposal and reflects the "
        "final cleaned site inventory."
    )
    document.add_paragraph(
        "Sites are classified as beaches, coastal waterfronts without a beach, lakefronts, or "
        "riverfronts for modelling. Two one-record hybrid categories are collapsed before regression: "
        "‘beach + coastal promenade’ is grouped with beaches, and ‘coastal promenade’ is grouped with "
        "coastal waterfronts without a beach. Beach is the reference group."
    )
    if (SOURCE_MEDIA / "image1.png").exists():
        add_picture(
            document,
            SOURCE_MEDIA / "image1.png",
            6.8,
            "Figure M1. Locations of the 114 blue–green waterfront sites. This site-location map is retained from the earlier working document; the updated accessibility maps will be prepared in ArcGIS.",
        )

    add_heading(document, "4.2.2 Multidimensional accessibility framework", 2)
    document.add_paragraph(
        "Accessibility is treated as a multidimensional construct with physical, visual, and haptic "
        "dimensions. Each dimension represents a distinct opportunity to reach, see, or potentially "
        "experience contact with water. The component outcomes are reported separately so that the "
        "aggregate score does not conceal which dimension drives an association."
    )
    if (SOURCE_MEDIA / "image2.jpg").exists():
        add_picture(
            document,
            SOURCE_MEDIA / "image2.jpg",
            5.6,
            "Figure M2. Conceptual multidimensional accessibility framework retained from the original working document.",
        )

    add_heading(document, "Physical accessibility", 3)
    document.add_paragraph(
        "Pedestrian access is the number of identified site access points per kilometre of the "
        "land-buffer boundary. It measures the density of ways to enter a site, not entrance quality, "
        "universal accessibility, or residential travel distance. Transit access is the number of bus "
        "stops inside the selected walk-time catchment. It measures nearby transit supply rather than "
        "service frequency, reliability, affordability, or a complete transit journey."
    )
    document.add_paragraph(
        "Pedestrian density and bus-stop counts are min–max scaled from 0 to 1 within each catchment. "
        "Combined physical accessibility is their equal arithmetic mean. Equal nominal weights do not "
        "guarantee equal empirical influence because the component variances differ; pedestrian and "
        "transit models are therefore shown alongside the combined score."
    )

    add_heading(document, "Visual accessibility", 3)
    document.add_paragraph(
        "Visual accessibility is the area of water visible from site observer points divided by the "
        "total water area in a common 1-km analysis area. Values slightly above one because of raster "
        "alignment are capped at one. Observers are generated hierarchically along paths using 100-m, "
        "50-m, and 25-m spacing, with 3–15 well-separated points per site and access-point fallback when "
        "necessary. A water cell is counted once if visible from at least one observer. The metric "
        "represents potential visible-water extent, not perceived scenic quality."
    )

    add_heading(document, "Haptic accessibility", 3)
    document.add_paragraph(
        "The primary Avery-style haptic measure is shoreline-contact length divided by the perimeter "
        "of the land buffer. It is best interpreted as shoreline exposure or potential contact rather "
        "than verified touchability: fences, vegetation, seawalls, water depth, safety, and other "
        "barriers are not fully observed. The earlier DTM–OSM shoreline-approachability proxy remains an "
        "exploratory construct-validity sensitivity measure. Its old 20-/30-minute coefficients are not "
        "carried into the present confirmatory tables because it has not yet been regenerated under the "
        "current 10-minute sample and common QA rules."
    )

    add_heading(document, "Multidimensional score", 3)
    document.add_paragraph(
        "Physical, visual, and haptic scores are each min–max scaled within a catchment and averaged with "
        "equal weight: multidimensional access = (scaled physical + scaled visual + scaled haptic) / 3. "
        "This aggregate is the synthesis outcome required by the approved proposal and the canonical "
        "accessibility input passed to Paper 4. The 10-minute score—not the earlier 20-minute score—is "
        "the current handoff."
    )

    add_heading(document, "4.2.3 Walking catchments and Census measures", 2)
    document.add_paragraph(
        "Network-based 5-, 10-, 20-, and 30-minute walking catchments are used to summarize surrounding "
        "Census characteristics and nearby bus stops. The 10-minute catchment is the main specification "
        "because it is a defensible neighbourhood walking scale and retains substantially more valid "
        "Census observations than five minutes. This choice was made conceptually and at the supervisor’s "
        "request before final reporting; it was not selected by searching across p-values. The other "
        "catchments are labelled sensitivity analyses."
    )
    document.add_paragraph(
        "Five Census predictors are evaluated: principal accommodation expenditure; percentage of "
        "residents below the low-income threshold; percentage with a bachelor’s degree or higher; "
        "percentage identifying as a visible minority; and percentage identifying as Indigenous. "
        "Outcome and predictor variables are standardized within each catchment-specific analytical "
        "sample, so reported coefficients are on a standard-deviation scale."
    )

    add_heading(document, "4.2.4 Data-quality exclusions and site controls", 2)
    document.add_paragraph(
        "The same eligibility rule is applied at every catchment: HasData must equal 1; all five Census "
        "predictors must be present; and no critical Census-quality flag may be present. Critical flags "
        "identify zero median-household-income values and related impossible or extreme combinations "
        "indicative of unreliable enrichment rather than genuine neighbourhood characteristics. This "
        "rule yields N = 90, 105, 112, and 114 at 5, 10, 20, and 30 minutes, respectively."
    )
    document.add_paragraph(
        "Controlled models add the four-category site type and the standardized natural logarithm of "
        "land-buffer area. The land buffer excludes adjacent water. Baseline models estimate the total "
        "observed association, while controlled models compare sites of similar type and terrestrial "
        "footprint. Because area and geometry contribute to several access constructions, the controlled "
        "model is a complementary estimand rather than automatically a more causal model."
    )

    add_heading(document, "4.2.5 Spatial representation and regression", 2)
    document.add_paragraph(
        "Each site is represented for neighbour calculations by an ArcGIS Feature To Point location "
        "constrained to lie inside its land-buffer polygon. These points are not catchment origins, "
        "access points, or unconstrained geometric centroids. Spatial weights are row-standardized "
        "k-nearest neighbours. Starting from k = 6, the smallest common k that connects every catchment "
        "sample is selected. The current data require k = 7 because the 30-minute graph has two components "
        "at k = 6."
    )
    document.add_paragraph(
        "For every outcome and catchment, the workflow fits standardized ordinary least squares (OLS) "
        "and tests residual spatial autocorrelation using a two-sided Moran’s I. OLS is retained when "
        "p ≥ .05. Otherwise, spatial autoregressive lag and error candidates are fit; among candidates "
        "that remove residual autocorrelation, the lowest-AIC model is selected. The controlled model "
        "begins with the baseline family to preserve comparability, but controlled OLS/SAR candidates "
        "are re-evaluated when residual Moran’s I remains significant."
    )
    document.add_paragraph(
        "OLS inference uses HC3 heteroskedasticity-robust standard errors. SAR models use model-based "
        "standard errors, and SAR-lag direct, indirect, and total impacts are retained in local diagnostic "
        "artifacts because structural lag coefficients are not total marginal effects. The 30 controlled "
        "socioeconomic tests are also evaluated using Benjamini–Hochberg false-discovery-rate adjustment. "
        "Variance-inflation factors are checked on the controlled design matrix; the maximum is 2.83."
    )

    add_heading(document, "4.2.6 Pairwise and extreme-site analyses", 2)
    document.add_paragraph(
        "Physical, visual, and haptic scores are compared across all 114 sites using Pearson and Spearman "
        "correlations, exact sample sizes, and vertical/horizontal median reference lines. Each ordered "
        "pair is also classified as high-high, high-low, low-high, or low-low using the pair-specific "
        "sample medians; values equal to a median are classified as high. These quadrants are descriptive "
        "groups, not statistical clusters. With 114 complete sites and no median ties, each dimension "
        "contains 57 high and 57 low sites; this balance mechanically makes HH equal LL and HL equal LH "
        "within each pair. Standardized dimension differences identify divergent sites "
        "for labelled scatterplots and spider profiles. "
        "A ternary plot displays each site’s relative physical–visual–haptic composition after normalizing "
        "the three non-negative scaled scores to shares; it describes composition, not total accessibility."
    )
    document.add_paragraph(
        "The pairwise correlation figure uses single-colour points, dashed median lines, and light "
        "horizontal background gridlines. The "
        "quadrant classifications are presented separately using quadrant-coloured points, dashed median "
        "lines, no background grid, and an in-panel count legend for each dimension pair."
    )
    document.add_paragraph(
        "The correlation/median, spider, and ternary elements requested in Asim’s shared file have been "
        "recalculated from the canonical data. Bland–Altman plots are not retained as a main result because "
        "they assess agreement between methods intended to measure the same quantity, whereas physical, "
        "visual, and haptic access are deliberately different constructs."
    )

    add_heading(document, "4.2.7 Protocol refinements relative to the proposal", 2)
    add_bullets(
        document,
        [
            "114 final sites replace the proposal’s preliminary 95/96-site inventory.",
            "Network walking catchments replace an adjacency-only Dissemination Area summary.",
            "Pedestrian entrance density and catchment transit supply are separated before being combined, improving construct clarity.",
            "Path-based 3–15 observer points and a union of visible water replace five random viewpoints and their average.",
            "An explicit OLS/SAR selection and residual-checking rule replaces a general statement that spatial models may be used.",
            "Component outcomes are reported beside the proposal’s aggregate, while multidimensional access remains the Paper 3 synthesis and Paper 4 input.",
        ],
    )
    add_note(
        document,
        "These are implementation refinements that must be reported transparently. The final manuscript "
        "should describe what was actually done rather than reproduce the proposal wording unchanged.",
    )


def add_results(
    document: Document,
    samples: list[dict[str, str]],
    regression: list[dict[str, str]],
    controls: list[dict[str, str]],
    correlations: list[dict[str, str]],
    quadrant_summary: list[dict[str, str]],
    fdr: list[dict[str, str]],
    diagnostics: list[dict[str, str]],
) -> None:
    add_heading(document, "4.3 Results", 1)
    add_heading(document, "4.3.1 Analytical sample and accessibility patterns", 2)
    document.add_paragraph(
        "All 114 sites have accessibility measurements. Census coverage and quality determine the "
        "regression sample: the main 10-minute analysis contains 105 sites after three critically flagged "
        "records are excluded. Larger catchments increase usable Census coverage, but those samples are "
        "used only to assess scale sensitivity."
    )
    sample_table = [
        [r["walktime_min"], r["total_sites"], r["HasData_equals_1"], r["critical_flags"], r["final_regression_N"]]
        for r in samples
    ]
    add_table(
        document,
        ["Catchment (min)", "Total sites", "With data", "Critical flags", "Final N"],
        sample_table,
        "Table 1. Sample sizes under the common Census-quality rule",
        font_size=8.5,
    )
    add_figure_placeholder(
        document,
        "INSERT UPDATED ARCGIS ACCESSIBILITY MAP(S) HERE",
        "Use the preferred original ArcGIS layout, replace the catchment background with the 10-minute catchment, and update the physical, visual, haptic, and multidimensional panels and legends as necessary.",
        "Figure 1. Physical, visual, haptic, and multidimensional accessibility at the 114 study sites using the current 10-minute specification. Maps show spatial pattern and should not be interpreted as neighbourhood causal effects.",
    )

    add_heading(document, "4.3.2 Relationships among access dimensions", 2)
    corr_by_pair = {row["pair"]: row for row in correlations}
    pv = corr_by_pair["Physical vs visual"]
    ph = corr_by_pair["Physical vs haptic"]
    vh = corr_by_pair["Visual vs haptic"]
    document.add_paragraph(
        f"Physical access is weakly related to visual access (Pearson r = {fmt_num(pv['Pearson_r'])}, "
        f"{fmt_p_statement(pv['Pearson_p'])}; Spearman ρ = {fmt_num(pv['Spearman_rho'])}, "
        f"{fmt_p_statement(pv['Spearman_p'])}) and haptic access (Pearson r = {fmt_num(ph['Pearson_r'])}, "
        f"{fmt_p_statement(ph['Pearson_p'])}; Spearman ρ = {fmt_num(ph['Spearman_rho'])}, "
        f"{fmt_p_statement(ph['Spearman_p'])}). Visual and haptic access have a modest positive relationship "
        f"(Pearson r = {fmt_num(vh['Pearson_r'])}, {fmt_p_statement(vh['Pearson_p'])}; "
        f"Spearman ρ = {fmt_num(vh['Spearman_rho'])}, {fmt_p_statement(vh['Spearman_p'])}). "
        "These patterns support treating the dimensions as related but "
        "non-interchangeable components of the aggregate."
    )
    corr_table = [
        [
            row["pair"],
            row["n"],
            fmt_num(row["Pearson_r"]),
            fmt_p(row["Pearson_p"]),
            fmt_num(row["Spearman_rho"]),
            fmt_p(row["Spearman_p"]),
        ]
        for row in correlations
    ]
    add_table(
        document,
        ["Pair", "n", "Pearson r", "p", "Spearman ρ", "p"],
        corr_table,
        "Table 2. Pairwise correlations among accessibility dimensions",
        font_size=8.3,
    )
    add_picture(
        document,
        FIGURE_DIR / "figure2_10min_pairwise_accessibility.png",
        7.0,
        "Figure 2. Pairwise accessibility scatterplots (n = 114). Pearson and Spearman statistics are shown for each dimension pair, and dashed lines show sample medians; the companion table identifies the largest standardized divergences.",
    )
    pair_order = ["Physical vs visual", "Physical vs haptic", "Visual vs haptic"]
    quadrant_order = {"HH": 0, "HL": 1, "LH": 2, "LL": 3}
    ordered_quadrants = sorted(
        quadrant_summary,
        key=lambda row: (
            pair_order.index(row["pair"]),
            quadrant_order[row["quadrant_code"]],
        ),
    )
    contrast_phrases = []
    for pair in pair_order:
        pair_rows = [row for row in ordered_quadrants if row["pair"] == pair]
        contrast_n = sum(
            int(row["site_count"])
            for row in pair_rows
            if row["quadrant_code"] in {"HL", "LH"}
        )
        pair_n = int(pair_rows[0]["pair_n"])
        contrast_phrases.append(
            f"{pair}: {contrast_n}/{pair_n} ({100 * contrast_n / pair_n:.1f}%)"
        )
    document.add_paragraph(
        "The median splits provide a transparent descriptive classification of cross-dimensional "
        "contrast. The combined high-low and low-high shares are "
        + "; ".join(contrast_phrases)
        + ". Because each dimension is split into 57 high and 57 low sites, paired quadrant counts are "
        "symmetric by construction. These proportions describe relative site profiles and do not test "
        "agreement or causality."
    )
    add_table(
        document,
        ["Pair", "Code", "Quadrant", "Sites", "% of pair"],
        [
            [
                row["pair"],
                row["quadrant_code"],
                row["quadrant"],
                row["site_count"],
                f"{float(row['percent']):.1f}",
            ]
            for row in ordered_quadrants
        ],
        "Table 2b. Median-based quadrant distribution for each accessibility-dimension pair",
        font_size=7.8,
    )
    add_picture(
        document,
        FIGURE_DIR / "figure2b_10min_quadrant_classification.png",
        6.9,
        "Figure 2b. Median-based quadrant classifications. Dashed lines show sample medians; in-panel legends report quadrant counts. Codes follow horizontal–vertical order.",
    )
    add_picture(
        document,
        FIGURE_DIR / "figure3_10min_extreme_site_profiles.png",
        6.7,
        "Figure 3. Physical, visual, and haptic profiles for the six most divergent sites. Profiles show component contrast, not statistical clusters.",
    )
    add_picture(
        document,
        FIGURE_DIR / "figure4_10min_accessibility_composition.png",
        6.6,
        "Figure 4. Relative physical–visual–haptic composition. Positions represent normalized shares and do not show the absolute magnitude of multidimensional access.",
    )

    add_heading(document, "4.3.3 Main 10-minute regression results", 2)
    ped_shelter = result_lookup(regression, "pedestrian", "z_shelter")
    ped_bach = result_lookup(regression, "pedestrian", "z_bachelors")
    tran_low = result_lookup(regression, "transit", "z_low_income")
    tran_shelter = result_lookup(regression, "transit", "z_shelter")
    phys_shelter = result_lookup(regression, "physical", "z_shelter")
    phys_low = result_lookup(regression, "physical", "z_low_income")
    river_multi = site_control_lookup(controls, "multidimensional", "site_typeriverfront")
    area_multi = site_control_lookup(controls, "multidimensional", "z_log_site_area")

    add_bullets(
        document,
        [
            f"Pedestrian access (SAR lag): accommodation expenditure is negatively associated after controls (β = {beta_from_cell(ped_shelter['controlled_beta_SE'])}, {fmt_p_statement(ped_shelter['controlled_p'])}). The positive baseline association with bachelor’s-degree share (β = {beta_from_cell(ped_bach['baseline_beta_SE'])}, {fmt_p_statement(ped_bach['baseline_p'])}) attenuates after site controls (β = {beta_from_cell(ped_bach['controlled_beta_SE'])}, {fmt_p_statement(ped_bach['controlled_p'])}).",
            f"Transit access (SAR error): low-income share is positively associated after controls (β = {beta_from_cell(tran_low['controlled_beta_SE'])}, {fmt_p_statement(tran_low['controlled_p'])}). Accommodation expenditure is negative (β = {beta_from_cell(tran_shelter['controlled_beta_SE'])}, {fmt_p_statement(tran_shelter['controlled_p'])}), but this result does not survive the 30-test false-discovery-rate adjustment.",
            f"Combined physical access (SAR error): accommodation expenditure is negative (β = {beta_from_cell(phys_shelter['controlled_beta_SE'])}, {fmt_p_statement(phys_shelter['controlled_p'])}) and low-income share is positive (β = {beta_from_cell(phys_low['controlled_beta_SE'])}, {fmt_p_statement(phys_low['controlled_p'])}). Both survive the false-discovery-rate adjustment.",
            "Visual access (SAR lag): no socioeconomic predictor is significant after adjustment. Relative to beaches, coastal waterfronts without a beach and riverfront sites have lower visual scores, while land-buffer area is positive.",
            "Haptic access (OLS with HC3 standard errors): no socioeconomic predictor or site-type contrast is significant. Land-buffer area has a strong positive association.",
            f"Multidimensional access (controlled SAR error): no socioeconomic predictor is significant. Riverfront sites are lower than beaches (β = {fmt_num(river_multi['estimate'])}, {fmt_p_statement(river_multi['p_value'])}), while land-buffer area is positive (β = {fmt_num(area_multi['estimate'])}, {fmt_p_statement(area_multi['p_value'])}).",
        ],
    )
    document.add_paragraph(
        "The low-income direction in the physical models is not a general measure of park advantage. It "
        "is concentrated in transit access: low-income share is associated with nearby bus-stop supply "
        "but not pedestrian entrance density, visual access, haptic access, or the final aggregate. The "
        "combined score also varies more strongly with its transit component than with its pedestrian "
        "component. Severe multicollinearity is unlikely to explain the directions (maximum VIF = 2.83)."
    )

    reg_rows = []
    for outcome in OUTCOME_ORDER:
        for term in TERM_ORDER:
            row = result_lookup(regression, outcome, term)
            reg_rows.append(
                [
                    OUTCOME_LABEL[outcome],
                    row["term_label"],
                    row["baseline_family"].replace("_", " "),
                    row["baseline_beta_SE"],
                    fmt_p(row["baseline_p"]),
                    row["controlled_family"].replace("_", " "),
                    row["controlled_beta_SE"],
                    fmt_p(row["controlled_p"]),
                ]
            )
    add_table(
        document,
        ["Outcome", "Predictor", "Baseline family", "β (SE)", "p", "Controlled family", "β (SE)", "p"],
        reg_rows,
        "Table 3. Complete 10-minute baseline and controlled socioeconomic regression results",
        font_size=6.8,
        first_column_shading=True,
    )
    add_caption(
        document,
        "Note. N = 105. Continuous outcomes and predictors are standardized. Controlled models include site type and standardized log land-buffer area. Asterisks in β (SE): * p < .05; ** p < .01; *** p < .001. SAR coefficients are structural coefficients; OLS uses HC3 standard errors.",
    )
    add_picture(
        document,
        FIGURE_DIR / "figure5_10min_controlled_coefficients.png",
        7.0,
        "Figure 5. Controlled socioeconomic coefficients and 95% confidence intervals for the six accessibility outcomes.",
    )

    add_heading(document, "4.3.4 Site controls and spatial diagnostics", 2)
    document.add_paragraph(
        "Land-buffer area is negatively associated with pedestrian access and positively associated with "
        "transit, visual, haptic, and multidimensional access. This is plausible because site size and "
        "geometry enter several accessibility constructions. Site-type contrasts are strongest for visual "
        "and multidimensional outcomes; they should be interpreted relative to beaches and conditional on "
        "the other predictors."
    )
    controls_rows = [
        [
            OUTCOME_LABEL[row["outcome"]],
            row["term_label"],
            row["family"].replace("_", " "),
            fmt_num(row["estimate"]),
            fmt_num(row["std_error"]),
            f"[{fmt_num(row['conf_low'])}, {fmt_num(row['conf_high'])}]",
            fmt_p(row["p_value"]),
        ]
        for row in controls
    ]
    add_table(
        document,
        ["Outcome", "Control", "Family", "β", "SE", "95% CI", "p"],
        controls_rows,
        "Table 4. Site-type and land-buffer-area coefficients in the controlled 10-minute models",
        font_size=7.1,
        first_column_shading=True,
    )

    controlled_diag = [row for row in diagnostics if row["specification"] == "controlled"]
    diag_rows = [
        [
            OUTCOME_LABEL[row["outcome"]],
            row["family"].replace("_", " "),
            row["N"],
            fmt_num(row["AIC"], 2),
            row["spatial_parameter"] if row["spatial_parameter"] not in ("", "NA") else "—",
            fmt_num(row["spatial_estimate"]),
            fmt_p(row["spatial_p"]),
            fmt_num(row["residual_Moran_I"]),
            fmt_p(row["residual_Moran_p"]),
        ]
        for row in controlled_diag
    ]
    add_table(
        document,
        ["Outcome", "Final family", "N", "AIC", "Spatial parameter", "Estimate", "p", "Residual Moran I", "p"],
        diag_rows,
        "Table 5. Final controlled-model diagnostics",
        font_size=7.0,
    )
    document.add_paragraph(
        "All six final controlled models have non-significant residual Moran’s I (p > .05). The controlled "
        "multidimensional OLS model initially retained spatial autocorrelation (p = .013); the reported "
        "SAR-error replacement clears that pattern (p = .818). SAR lag and error candidates differ by "
        "less than two AIC units for the main spatial outcomes, so the selected family implements a "
        "reproducible rule but does not prove a specific spatial-generating mechanism."
    )

    add_heading(document, "4.3.5 Multiple-testing assessment", 2)
    survivors = [row for row in fdr if row["survives_BH_0_05"].upper() == "TRUE"]
    fdr_rows = [
        [
            OUTCOME_LABEL[row["outcome"]],
            row["term_label"],
            fmt_num(row["estimate"]),
            fmt_p(row["p_value"]),
            fmt_p(row["p_BH_across_main_tests"]),
        ]
        for row in survivors
    ]
    add_table(
        document,
        ["Outcome", "Predictor", "β", "Raw p", "BH q"],
        fdr_rows,
        "Table 6. Controlled socioeconomic associations surviving BH correction across 30 main tests",
        font_size=8.0,
    )
    document.add_paragraph(
        "Four associations remain below a 5% false-discovery-rate threshold: accommodation expenditure "
        "with pedestrian access; low-income share with transit access; and both variables with combined "
        "physical access. The nominal transit/accommodation association does not survive this correction."
    )

    add_heading(document, "4.3.6 Catchment sensitivity", 2)
    document.add_paragraph(
        "The accommodation–pedestrian association is negative at 5, 10, 20, and 30 minutes. The positive "
        "low-income–transit association appears at 10, 20, and 30 minutes. No socioeconomic predictor is "
        "significant for visual, haptic, or multidimensional access at any catchment. Other isolated "
        "p < .05 estimates vary with scale and are treated as exploratory rather than independent "
        "confirmatory findings. Appendix B reports every controlled sensitivity coefficient."
    )

    add_heading(document, "4.3.7 Synthesis for Paper 4", 2)
    document.add_paragraph(
        "The multidimensional accessibility score remains scientifically coherent as a synthesis input "
        "because its components are empirically distinct and because none alone represents overall "
        "access. For Paper 4, the correct current handoff is the 10-minute score for all 114 sites. The "
        "regression sample is smaller (N = 105) only because socioeconomic models require valid Census "
        "predictors; this does not require discarding accessibility scores for the other nine sites from "
        "the Mapbox perception analysis. Paper 4 outputs previously derived from the old 20-minute handoff "
        "must be labelled provisional until regenerated."
    )


def make_landscape_section(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    add_page_number(section.footer.paragraphs[0])


def add_appendices(
    document: Document,
    analysis_rows: list[dict[str, str]],
    sensitivity: list[dict[str, str]],
) -> None:
    make_landscape_section(document)
    add_heading(document, "Appendix A. Study-site inventory", 1)
    ten_min = sorted(
        [row for row in analysis_rows if row["walktime_min"] == "10"],
        key=lambda row: int(row["park_num"]),
    )
    site_rows = [
        [
            row["park_num"],
            row["site_name"],
            row["municipality"],
            row["site_type"],
            "Yes" if row["regression_eligible"].upper() == "TRUE" else "No",
            row["census_exclusion_reason"] if row["census_exclusion_reason"] not in ("", "NA") else "—",
        ]
        for row in ten_min
    ]
    add_table(
        document,
        ["Site ID", "Site name", "Municipality / jurisdiction", "Site type", "10-min eligible", "Exclusion reason"],
        site_rows,
        "Table A1. Final 114-site inventory and 10-minute regression eligibility",
        font_size=6.6,
    )

    add_heading(document, "Appendix B. Complete catchment sensitivity results", 1)
    sensitivity_sorted = sorted(
        sensitivity,
        key=lambda row: (
            int(row["catchment_min"]),
            OUTCOME_ORDER.index(row["outcome"]),
            TERM_ORDER.index(row["term"]),
        ),
    )
    sensitivity_rows = [
        [
            row["catchment_min"],
            OUTCOME_LABEL[row["outcome"]],
            row["term_label"],
            row["family"].replace("_", " "),
            row["N"],
            fmt_num(row["estimate"]),
            fmt_num(row["std_error"]),
            f"[{fmt_num(row['conf_low'])}, {fmt_num(row['conf_high'])}]",
            fmt_p(row["p_value"]),
        ]
        for row in sensitivity_sorted
    ]
    add_table(
        document,
        ["Min", "Outcome", "Predictor", "Family", "N", "β", "SE", "95% CI", "p"],
        sensitivity_rows,
        "Table B1. Controlled socioeconomic coefficients at 5, 10, 20, and 30 minutes",
        font_size=6.2,
        first_column_shading=True,
    )
    add_caption(
        document,
        "Note. The 10-minute specification is primary. Other catchments assess sensitivity to spatial scale and should not be used to select whichever result is most significant.",
    )

    add_heading(document, "Appendix C. Reporting and reproducibility notes", 1)
    add_bullets(
        document,
        [
            "Public tables and PNG figures are stored under outputs/accessibility_analysis and may be committed to GitHub.",
            "Raw candidate-model diagnostics, extracted source-document media, and this Word deliverable are local-only and Git-ignored.",
            "All data directories are local-only. The repository records required inputs and workflow order but does not commit the large or restricted datasets.",
            "The canonical model entry point is scripts/accessibility_analysis/run_all.R; it regenerates public Paper 3 tables and PNG figures.",
            "The ArcGIS accessibility map is inserted manually at the marked Figure 1 placeholder.",
            "The Word document is generated after analysis by scripts/accessibility_analysis/03_build_methods_results_docx.py and is intentionally not called by the public R workflow.",
        ],
    )

    add_heading(document, "Reference retained from the working draft", 1)
    document.add_paragraph(
        "Accessibility Standards Canada. (2023, June). 6. Common accessibility measures. "
        "CAN-ASC-2.1 – Outdoor Spaces. https://accessible.canada.ca/standards-and-technical-guides/"
        "standards-and-technical-guides-database/can-asc-21-outdoor-spaces/6-common-accessibility-measures"
    )
    add_note(
        document,
        "Before journal submission, reconcile all author–year citations from the Introduction and full "
        "manuscript with the reference manager. The source Methods/Results draft contained only the "
        "Accessibility Standards Canada reference in its reference section.",
    )


def build_document() -> Path:
    output_file = requested_output_file()
    required = [
        DATA_FILE,
        TABLE_DIR / "table1_sample_sizes_by_catchment.csv",
        TABLE_DIR / "table2_10min_full_regression.csv",
        TABLE_DIR / "table3_10min_site_controls.csv",
        TABLE_DIR / "table4_controlled_catchment_sensitivity.csv",
        TABLE_DIR / "table5_10min_pairwise_correlations.csv",
        TABLE_DIR / "table5b_10min_quadrant_summary.csv",
        TABLE_DIR / "table6_10min_multiple_testing.csv",
        TABLE_DIR / "table7_10min_model_diagnostics.csv",
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing canonical input(s):\n" + "\n".join(missing))

    analysis = read_csv(DATA_FILE)
    samples = read_csv(TABLE_DIR / "table1_sample_sizes_by_catchment.csv")
    regression = read_csv(TABLE_DIR / "table2_10min_full_regression.csv")
    controls = read_csv(TABLE_DIR / "table3_10min_site_controls.csv")
    sensitivity = read_csv(TABLE_DIR / "table4_controlled_catchment_sensitivity.csv")
    correlations = read_csv(TABLE_DIR / "table5_10min_pairwise_correlations.csv")
    quadrant_summary = read_csv(TABLE_DIR / "table5b_10min_quadrant_summary.csv")
    fdr = read_csv(TABLE_DIR / "table6_10min_multiple_testing.csv")
    diagnostics = read_csv(TABLE_DIR / "table7_10min_model_diagnostics.csv")

    document = Document()
    configure_document(document)
    document.core_properties.title = "Paper 3 Methods and Results 2.0"
    document.core_properties.subject = "Canonical 10-minute Paper 3 analysis"
    document.core_properties.author = "Urban Blue–Green Accessibility Project"
    document.core_properties.comments = "Private working document generated from canonical Paper 3 outputs."

    add_title_page(document)
    add_executive_summary(document, samples)
    add_methods(document, analysis)
    add_results(
        document, samples, regression, controls, correlations,
        quadrant_summary, fdr, diagnostics
    )
    add_appendices(document, analysis, sensitivity)

    output_file.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_file)
    return output_file


if __name__ == "__main__":
    result = build_document()
    print(f"Created {result}")
